from docx import Document
import os
from docx.text.paragraph import Paragraph
from docx.oxml.text.run import CT_R

class DocxTableConverter:
    def __init__(self):
        self.markdown_tables = []

    def get_formatted_text(self, paragraph):
        """Extract text with formatting from a paragraph"""
        text_parts = []
        
        for run in paragraph.runs:
            text = run.text.strip()
            if text:
                # Handle bold text
                if run.bold:
                    text = f"**{text}**"
                text_parts.append(text)
        
        return ' '.join(text_parts)

    def process_cell_content(self, cell):
        """Process cell content preserving formatting and bullets"""
        lines = []
        
        for paragraph in cell.paragraphs:
            # Skip empty paragraphs
            if not paragraph.text.strip():
                continue
                
            line = ""
            # Check if paragraph is a bullet point
            if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
                line = "* "
            
            # Add formatted text
            line += self.get_formatted_text(paragraph)
            
            if line:
                lines.append(line)
        
        return '<br>'.join(lines) if lines else ''

    def convert_table_to_markdown(self, table):
        """Convert a single table to markdown format with preserved formatting"""
        markdown_rows = []
        
        # Get headers
        headers = []
        for cell in table.rows[0].cells:
            headers.append(self.process_cell_content(cell))
        markdown_rows.append('| ' + ' | '.join(headers) + ' |')
        
        # Add separator row with alignment
        # Using HTML style alignment to force top alignment
        column_count = len(headers)
        style_row = '<style>\n  td, th {\n    vertical-align: top !important;\n  }\n</style>\n'
        markdown_rows.insert(0, style_row)
        
        # Add standard separator row
        separator = '|' + '|'.join(['---' for _ in headers]) + '|'
        markdown_rows.append(separator)
        
        # Get data rows
        for row in table.rows[1:]:
            cells = []
            for cell in row.cells:
                cell_content = self.process_cell_content(cell)
                cells.append(cell_content)
            markdown_rows.append('| ' + ' | '.join(cells) + ' |')
        
        return '\n'.join(markdown_rows)

    def convert_file(self, input_file, output_file):
        """Convert all tables in a DOCX file to markdown format"""
        try:
            # Check if input file exists
            if not os.path.exists(input_file):
                raise FileNotFoundError(f"Input file not found: {input_file}")

            # Load the document
            print(f"Converting tables from: {input_file}")
            doc = Document(input_file)
            
            # Convert each table
            for i, table in enumerate(doc.tables, 1):
                print(f"Converting table {i}...")
                markdown_table = self.convert_table_to_markdown(table)
                self.markdown_tables.append(f"\n### Table {i}\n\n{markdown_table}\n")
            
            # Write to output file
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write("# Converted Tables\n\n")
                f.write('\n'.join(self.markdown_tables))
            
            print(f"\nConversion completed! Output saved to: {output_file}")
            print(f"Found and converted {len(self.markdown_tables)} tables")

        except Exception as e:
            print(f"Error during conversion: {str(e)}")

def main():
    try:
        converter = DocxTableConverter()
        
        # Example usage
        input_file = "input.docx"  # Replace with your DOCX file
        output_file = "output.md"  # Replace with desired output path
        
        converter.convert_file(input_file, output_file)
        
    except Exception as e:
        print(f"\nError: {str(e)}")

if __name__ == "__main__":
    main() 